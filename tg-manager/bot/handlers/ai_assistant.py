"""AI-ассистент через OpenRouter — изолирован по user_id, доступ только к данным пользователя.

READ-инструменты возвращают данные для анализа.
ACTION-инструменты запрашивают подтверждение у пользователя перед выполнением.
"""

from __future__ import annotations
import asyncio
import json
import logging
import os
import re
from html import escape
from io import BytesIO
from pathlib import Path
import asyncpg
import aiohttp
from aiogram import Router, F
from aiogram.filters import Command
from aiogram.fsm.context import FSMContext
from aiogram.types import CallbackQuery, Message
from aiogram.utils.keyboard import InlineKeyboardBuilder
from bot.callbacks import AiCb, BmCb
from bot.states import AiChat
from bot.utils.subscription import require_plan
from bot.utils.ai_tools import TOOL_DEFINITIONS, run_tool, execute_action
from bot.utils.op_helpers import safe_answer

from services import ai_memory
from services.ai_providers import configured_providers
from services.logger import log_exc_swallow

# Регулярное выражение для тега [MEMORY: заголовок | тело]
_MEMORY_TAG_RE = re.compile(
    r"\[MEMORY:\s*([^\|\]]{1,180})\|([^\]]{1,2000})\]",
    re.DOTALL,
)

router = Router()
log = logging.getLogger(__name__)

_SYSTEM_PROMPT = (
    "Ты AI-ассистент платформы Infragram OS — корпоративной системы управления Telegram-инфраструктурой.\n"
    "Ты можешь АНАЛИЗИРОВАТЬ данные пользователя И РЕАЛЬНО ВЫПОЛНЯТЬ задачи управления.\n\n"
    "ВОЗМОЖНОСТИ (инструменты):\n"
    "READ — данные:\n"
    "- get_my_bots: список ботов + статистика\n"
    "- get_bot_details: детали конкретного бота\n"
    "- get_network_stats: сводная статистика сети\n"
    "- get_audience_activity: сегментация аудитории (hot/warm/cold/lost)\n"
    "- get_growth_trend: динамика роста за N дней\n"
    "- get_seo_recommendations: SEO-оценка бота\n"
    "- get_my_accounts: список Telegram-аккаунтов\n"
    "- get_my_channels: список каналов из кэша\n"
    "- get_broadcast_history: история рассылок\n\n"
    "ACTION — действия (требуют подтверждения пользователя):\n"
    "- create_channel: СОЗДАТЬ канал в Telegram через подключённый аккаунт\n"
    "- create_group: СОЗДАТЬ группу/супергруппу в Telegram\n"
    "- create_bot: СОЗДАТЬ бота через BotFather\n"
    "- bulk_create_channels: МАССОВОЕ создание каналов (3-50 шт) с умными задержками — ставится в очередь\n"
    "- post_to_channel: опубликовать пост в канал\n"
    "- launch_broadcast: запустить рассылку боту сейчас\n"
    "- schedule_broadcast: запланировать рассылку на через N минут\n"
    "- update_bot_profile: обновить имя/описание бота\n\n"
    "КОГДА ЧТО ИСПОЛЬЗОВАТЬ:\n"
    "- create_channel — для ОДНОГО канала\n"
    "- create_group — для ОДНОЙ группы\n"
    "- create_bot — для ОДНОГО бота\n"
    "- bulk_create_channels — когда пользователь просит СОЗДАТЬ НЕСКОЛЬКО каналов (3+)\n\n"
    "ПРОТОКОЛ ДЕЙСТВИЙ:\n"
    "1. Когда пользователь просит создать канал/группу/бота/сеть присутствия — используй соответствующий инструмент\n"
    "2. Для массового создания (>2 объектов) — используй bulk_create_channels\n"
    "3. Действие НЕ выполнится сразу — пользователь подтвердит его\n"
    "4. После подготовки действия — сообщи что подготовлено и ждёт подтверждения\n"
    "5. НЕ ГОВОРИ что не можешь создать присутствие — ты УМЕЕШЬ создавать каналы, группы и ботов\n\n"
    "ПАМЯТЬ:\n"
    "Если в ходе диалога ты обнаруживаешь важную информацию о пользователе, его бизнесе,\n"
    "предпочтениях или целях — которую стоит запомнить для будущих сессий — добавь в конец ответа тег:\n"
    "[MEMORY: заголовок | подробное описание для сохранения]\n"
    "Пример: [MEMORY: Основной продукт | Пользователь продаёт курсы по Python, целевая аудитория — новички]\n"
    "Используй этот тег только для действительно важной информации. Можно добавить один тег за ответ.\n\n"
    "ВАЖНО:\n"
    "- Отвечай только на русском языке\n"
    "- Давай конкретные советы с числами\n"
    "- Всегда предлагай реальные действия, а не только рекомендации\n"
    "- Доступ только к данным текущего пользователя"
)

_MAX_TURNS = 100

# ── Fast action patterns (no LLM needed) ──────────────────────────────────────

_FAST_ACTION_PATTERNS: list[tuple] = [
    # "создай 5 каналов <prefix>" — bulk first (more specific)
    (
        re.compile(
            r"созда(?:й|йте|ть|вай)\s+(\d+)\s+канал\w*\s+(.+)", re.I | re.DOTALL
        ),
        "bulk_channels",
    ),
    # "создай канал <title>"
    (
        re.compile(r"созда(?:й|йте|ть|вай)\s+канал(?!\w)\s+(.+)", re.I | re.DOTALL),
        "channel",
    ),
    # "создай группу <title>"
    (
        re.compile(
            r"созда(?:й|йте|ть|вай)\s+(?:чат|групп\w+)\s+(.+)", re.I | re.DOTALL
        ),
        "group",
    ),
    # "создай бота <name>"
    (
        re.compile(r"созда(?:й|йте|ть|вай)\s+бот\w*\s+(.+)", re.I | re.DOTALL),
        "bot",
    ),
]


async def _fast_parse_action(
    text: str, pool: asyncpg.Pool, user_id: int
) -> dict | None:
    """
    Detect simple create/action intents without LLM.
    Returns {"__pending__": True, "action_data": ..., "ai_message": ...} or None.
    """
    from bot.utils.ai_tools import (
        action_create_channel,
        action_create_group,
        action_create_bot,
        action_bulk_create_channels,
    )

    stripped = text.strip()

    for pattern, kind in _FAST_ACTION_PATTERNS:
        m = pattern.search(stripped)
        if not m:
            continue

        if kind == "bulk_channels":
            try:
                count = min(50, max(1, int(m.group(1))))
            except ValueError:
                continue
            prefix = m.group(2).strip().strip("«»\"' \t\n")[:80]
            if len(prefix) < 2:
                continue
            result = await action_bulk_create_channels(
                pool, user_id, prefix=prefix, count=count
            )
            if "error" in result:
                return None
            return {
                "__pending__": True,
                "action_data": result,
                "ai_message": (
                    f"⚡ Подготовил операцию: создать <b>{count}</b> каналов "
                    f"с префиксом «{prefix}»."
                ),
            }

        elif kind == "channel":
            title = m.group(1).strip().strip("«»\"' \t\n")[:80]
            if len(title) < 2:
                continue
            result = await action_create_channel(pool, user_id, title=title)
            if "error" in result:
                return None
            return {
                "__pending__": True,
                "action_data": result,
                "ai_message": f"⚡ Подготовил создание канала «{title}».",
            }

        elif kind == "group":
            title = m.group(1).strip().strip("«»\"' \t\n")[:80]
            if len(title) < 2:
                continue
            result = await action_create_group(pool, user_id, title=title)
            if "error" in result:
                return None
            return {
                "__pending__": True,
                "action_data": result,
                "ai_message": f"⚡ Подготовил создание группы «{title}».",
            }

        elif kind == "bot":
            name = m.group(1).strip().strip("«»\"' \t\n")[:80]
            if len(name) < 2:
                continue
            result = await action_create_bot(pool, user_id, name=name)
            if "error" in result:
                return None
            return {
                "__pending__": True,
                "action_data": result,
                "ai_message": f"⚡ Подготовил создание бота «{name}».",
            }

    return None


_DEFAULT_MAX_FILE_BYTES = 1_048_576
_DEFAULT_MAX_FILE_CHARS = 60_000
_TELEGRAM_TEXT_LIMIT = 3900
_SUPPORTED_TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".json",
    ".jsonl",
    ".csv",
    ".tsv",
    ".yaml",
    ".yml",
    ".xml",
    ".html",
    ".htm",
    ".css",
    ".scss",
    ".js",
    ".jsx",
    ".ts",
    ".tsx",
    ".py",
    ".sql",
    ".toml",
    ".ini",
    ".env",
    ".log",
    ".prisma",
    ".sh",
    ".ps1",
}
_SUPPORTED_STRUCTURED_EXTENSIONS = {".pdf", ".docx", ".xlsx"}
_SUPPORTED_FILE_HINT = "txt, md, csv, json, yaml, код, log, pdf, docx, xlsx"


def _get_positive_int_env(name: str, default: int) -> int:
    raw = os.getenv(name)
    if not raw:
        return default
    try:
        value = int(raw)
    except ValueError:
        log.warning("Invalid %s=%r, using default %s", name, raw, default)
        return default
    return max(1024, value)


_MAX_FILE_BYTES = _get_positive_int_env("AI_FILE_MAX_BYTES", _DEFAULT_MAX_FILE_BYTES)
_MAX_FILE_CHARS = _get_positive_int_env("AI_FILE_MAX_CHARS", _DEFAULT_MAX_FILE_CHARS)


def _openai_tools() -> list:
    result = []
    for t in TOOL_DEFINITIONS:
        result.append(
            {
                "type": "function",
                "function": {
                    "name": t["name"],
                    "description": t["description"],
                    "parameters": t["input_schema"],
                },
            }
        )
    return result


def _truncate_file_text(text: str) -> str:
    text = text.replace("\x00", "").strip()
    if len(text) <= _MAX_FILE_CHARS:
        return text
    return (
        text[:_MAX_FILE_CHARS].rstrip()
        + "\n\n[Файл длинный, поэтому дальше текст обрезан. Пришлите нужную часть отдельно, если нужны детали ниже.]"
    )


def _decode_text_bytes(raw: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1251", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _extract_pdf_text(raw: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("Для чтения PDF нужна библиотека pypdf.") from exc

    reader = PdfReader(BytesIO(raw))
    parts: list[str] = []
    max_pages = 20
    for index in range(min(len(reader.pages), max_pages)):
        page_number = index + 1
        page = reader.pages[index]
        page_text = page.extract_text() or ""
        if page_text.strip():
            parts.append(f"[Страница {page_number}]\n{page_text.strip()}")
    if len(reader.pages) > max_pages:
        parts.append(f"[Показаны первые {max_pages} страниц из {len(reader.pages)}.]")
    return "\n\n".join(parts)


def _extract_docx_text(raw: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("Для чтения DOCX нужна библиотека python-docx.") from exc

    doc = Document(BytesIO(raw))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table_index, table in enumerate(doc.tables, start=1):
        rows: list[str] = []
        for row in table.rows[:100]:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append(f"[Таблица {table_index}]\n" + "\n".join(rows))
    return "\n\n".join(parts)


def _extract_xlsx_text(raw: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("Для чтения XLSX нужна библиотека openpyxl.") from exc

    workbook = load_workbook(BytesIO(raw), read_only=True, data_only=True)
    parts: list[str] = []
    max_sheets = 5
    max_rows = 120
    max_columns = 20
    try:
        for sheet in workbook.worksheets[:max_sheets]:
            rows: list[str] = []
            for row in sheet.iter_rows(
                max_row=max_rows, max_col=max_columns, values_only=True
            ):
                values = ["" if value is None else str(value) for value in row]
                if any(value.strip() for value in values):
                    rows.append("\t".join(values).rstrip())
            if rows:
                parts.append(f"[Лист: {sheet.title}]\n" + "\n".join(rows))
        if len(workbook.worksheets) > max_sheets:
            parts.append(
                f"[Показаны первые {max_sheets} листов из {len(workbook.worksheets)}.]"
            )
    finally:
        workbook.close()
    return "\n\n".join(parts)


def _extract_document_text(file_name: str, mime_type: str | None, raw: bytes) -> str:
    suffix = Path(file_name).suffix.lower()
    mime_type = (mime_type or "").lower()

    if suffix in _SUPPORTED_TEXT_EXTENSIONS or mime_type.startswith("text/"):
        return _decode_text_bytes(raw)
    if suffix == ".pdf":
        return _extract_pdf_text(raw)
    if suffix == ".docx":
        return _extract_docx_text(raw)
    if suffix == ".xlsx":
        return _extract_xlsx_text(raw)

    supported = sorted(_SUPPORTED_TEXT_EXTENSIONS | _SUPPORTED_STRUCTURED_EXTENSIONS)
    raise ValueError(
        f"Формат пока не поддерживается. Поддерживаются: {', '.join(supported)}"
    )


async def _read_attached_document(message: Message) -> tuple[str, str]:
    document = message.document
    if not document:
        raise ValueError("Файл не найден в сообщении.")

    file_name = document.file_name or "file"
    file_size = document.file_size or 0
    if file_size > _MAX_FILE_BYTES:
        size_mb = _MAX_FILE_BYTES / 1024 / 1024
        raise ValueError(
            f"Файл слишком большой. Сейчас можно отправлять файлы до {size_mb:.1f} МБ."
        )

    if message.bot is None:
        raise RuntimeError(
            "Бот не готов скачать файл. Попробуйте отправить файл еще раз."
        )

    bot_file = await message.bot.get_file(document.file_id)
    downloaded = await message.bot.download_file(bot_file.file_path)
    if hasattr(downloaded, "seek"):
        downloaded.seek(0)
    raw = downloaded.read() if hasattr(downloaded, "read") else bytes(downloaded or b"")
    if len(raw) > _MAX_FILE_BYTES:
        size_mb = _MAX_FILE_BYTES / 1024 / 1024
        raise ValueError(
            f"Файл слишком большой. Сейчас можно отправлять файлы до {size_mb:.1f} МБ."
        )

    text = _extract_document_text(file_name, document.mime_type, raw)
    text = _truncate_file_text(text)
    if not text:
        raise ValueError("В файле не удалось найти читаемый текст.")
    return file_name, text


def _build_file_prompt(file_name: str, file_text: str, caption: str | None) -> str:
    task = (
        caption.strip()
        if caption and caption.strip()
        else "Изучи файл и помоги по его содержимому."
    )
    return (
        f"{task}\n\n"
        f"Пользователь отправил файл: {file_name}\n"
        "Проанализируй содержимое файла. Если в файле есть ошибки, риски или важные выводы, объясни их простыми словами.\n\n"
        f"--- НАЧАЛО ФАЙЛА: {file_name} ---\n"
        f"{file_text}\n"
        f"--- КОНЕЦ ФАЙЛА: {file_name} ---"
    )


def _split_telegram_text(text: str) -> list[str]:
    if len(text) <= _TELEGRAM_TEXT_LIMIT:
        return [text]

    chunks: list[str] = []
    remaining = text
    while remaining:
        if len(remaining) <= _TELEGRAM_TEXT_LIMIT:
            chunks.append(remaining)
            break

        split_at = max(
            remaining.rfind("\n\n", 0, _TELEGRAM_TEXT_LIMIT),
            remaining.rfind("\n", 0, _TELEGRAM_TEXT_LIMIT),
            remaining.rfind(" ", 0, _TELEGRAM_TEXT_LIMIT),
        )
        if split_at < _TELEGRAM_TEXT_LIMIT // 2:
            split_at = _TELEGRAM_TEXT_LIMIT

        chunks.append(remaining[:split_at].rstrip())
        remaining = remaining[split_at:].lstrip()

    return [chunk for chunk in chunks if chunk]


async def _edit_or_answer_long(
    thinking_message: Message,
    source_message: Message,
    text: str,
    reply_markup=None,
) -> None:
    chunks = _split_telegram_text(text)
    if len(chunks) == 1:
        try:
            await thinking_message.edit_text(
                chunks[0], parse_mode="HTML", reply_markup=reply_markup
            )
            return
        except Exception:
            try:
                await thinking_message.edit_text(
                    chunks[0], parse_mode=None, reply_markup=reply_markup
                )
                return
            except Exception:
                await source_message.answer(
                    chunks[0], parse_mode=None, reply_markup=reply_markup
                )
                return

    try:
        await thinking_message.edit_text(chunks[0], parse_mode=None)
    except Exception:
        await source_message.answer(chunks[0], parse_mode=None)

    for chunk in chunks[1:-1]:
        await source_message.answer(chunk, parse_mode=None)

    await source_message.answer(chunks[-1], parse_mode=None, reply_markup=reply_markup)


async def _call_ai_providers(
    messages: list,
    pool: asyncpg.Pool,
    user_id: int,
    http: aiohttp.ClientSession | None = None,
) -> str | dict:
    """OpenAI-compatible provider failover with Infragram memory context."""
    providers = configured_providers()
    if not providers:
        return (
            "⚠️ <b>AI-ассистент не настроен</b>\n\n"
            "Добавьте хотя бы один ключ: <code>OPENROUTER_API_KEY</code>, "
            "<code>GROQ_API_KEY</code> или <code>GEMINI_API_KEY</code>."
        )

    try:
        from openai import AsyncOpenAI
    except ImportError:
        return "⚠️ Библиотека openai не установлена."

    memory_context = ""
    try:
        last_user_message = next(
            (
                m.get("content", "")
                for m in reversed(messages)
                if m.get("role") == "user"
            ),
            "",
        )
        memory_context = ai_memory.format_for_prompt(
            await ai_memory.search(pool, user_id, str(last_user_message), limit=8)
        )
    except Exception:
        log_exc_swallow(log, "Не удалось загрузить память AI")

    base_messages = [{"role": "system", "content": _SYSTEM_PROMPT}]
    if memory_context:
        base_messages.append({"role": "system", "content": memory_context})
    base_messages.extend(list(messages))

    tools = _openai_tools()
    primary_label = f"{providers[0].name}/{providers[0].models[0]}"
    attempts = [
        (provider, model, max_tokens, use_tools)
        for use_tools in (True, False)
        for max_tokens in (1500, 600)
        for provider in providers
        for model in provider.models
    ]
    seen: set[tuple[str, str, int, bool]] = set()

    for provider, model, max_tokens, use_tools in attempts:
        attempt_key = (provider.name, model, max_tokens, use_tools)
        if attempt_key in seen:
            continue
        seen.add(attempt_key)
        client = AsyncOpenAI(
            api_key=provider.api_key,
            base_url=provider.base_url,
            timeout=25.0,
        )
        current_messages = list(base_messages)
        try:
            pending_action_data = None
            for _ in range(8):
                completion_kwargs = {
                    "model": model,
                    "messages": current_messages,
                    "max_tokens": max_tokens,
                }
                if use_tools:
                    completion_kwargs["tools"] = tools
                    completion_kwargs["tool_choice"] = "auto"
                response = await client.chat.completions.create(**completion_kwargs)
                choice = response.choices[0]
                msg = choice.message

                if choice.finish_reason == "tool_calls" and msg.tool_calls:
                    assistant_msg: dict = {
                        "role": "assistant",
                        "content": msg.content or "",
                        "tool_calls": [
                            {
                                "id": tc.id,
                                "type": "function",
                                "function": {
                                    "name": tc.function.name,
                                    "arguments": tc.function.arguments,
                                },
                            }
                            for tc in msg.tool_calls
                        ],
                    }
                    current_messages.append(assistant_msg)

                    for tc in msg.tool_calls:
                        try:
                            args = json.loads(tc.function.arguments)
                        except Exception:
                            args = {}
                        result_text = await run_tool(
                            tc.function.name, args, pool, user_id, http
                        )

                        try:
                            parsed = json.loads(result_text)
                            if isinstance(parsed, dict) and "pending_action" in parsed:
                                pending_action_data = parsed
                                result_text = json.dumps(
                                    {
                                        "status": "pending_confirmation",
                                        "message": f"Действие «{parsed['pending_action']}» ждёт подтверждения пользователя.",
                                        "preview": parsed.get("preview", ""),
                                    },
                                    ensure_ascii=False,
                                )
                        except Exception:
                            log_exc_swallow(log, "Не удалось распарсить tool JSON")

                        current_messages.append(
                            {
                                "role": "tool",
                                "tool_call_id": tc.id,
                                "content": result_text,
                            }
                        )
                    continue

                ai_reply = msg.content or "Нет ответа."
                label = f"{provider.name}/{model}"
                suffix = f"\n\n<i>Модель: {label}</i>" if label != primary_label else ""
                if not use_tools:
                    suffix += "\n<i>Режим: без инструментов</i>"
                full_reply = ai_reply + suffix
                if pending_action_data:
                    return {
                        "__pending__": True,
                        "action_data": pending_action_data,
                        "ai_message": full_reply,
                    }
                return full_reply

            return "Превышен лимит итераций инструментов."

        except Exception as e:
            err_str = str(e).lower()
            should_retry = any(
                x in err_str
                for x in (
                    "rate",
                    "limit",
                    "429",
                    "quota",
                    "overloaded",
                    "timeout",
                    "model_not_found",
                    "not found",
                    "404",
                    "invalid model",
                    "402",
                    "credits",
                    "payment",
                    "billing",
                    "insufficient",
                    "tool",
                )
            )
            if should_retry:
                log.warning(
                    "AI provider %s model %s failed (%s), trying next",
                    provider.name,
                    model,
                    type(e).__name__,
                )
                await asyncio.sleep(
                    1.5 if "429" in err_str or "rate" in err_str else 0.4
                )
                continue
            log.exception("AI provider %s model %s error: %s", provider.name, model, e)
            return (
                f"⚠️ Ошибка AI-ассистента ({provider.name}/{model}): "
                f"{type(e).__name__}: {str(e)[:120]}"
            )

    return (
        "⚠️ <b>AI-ассистент временно недоступен</b>\n\n"
        "Все настроенные модели не ответили. Проверьте ключи и доступность "
        "<code>OPENROUTER_API_KEY</code>, <code>GROQ_API_KEY</code>, "
        "<code>GEMINI_API_KEY</code>."
    )


async def _parse_and_save_memory(
    text: str,
    pool: asyncpg.Pool,
    user_id: int,
) -> tuple[str, int]:
    """Найти теги [MEMORY: ...] в тексте, сохранить их и вернуть очищенный текст + кол-во сохранённых."""
    saved = 0
    clean = text
    for match in _MEMORY_TAG_RE.finditer(text):
        title = match.group(1).strip()
        body = match.group(2).strip()
        if not body:
            continue
        try:
            await ai_memory.remember(
                pool,
                user_id,
                body,
                title=title or body[:80],
                kind="insight",
                source="ai",
            )
            saved += 1
        except Exception:
            log_exc_swallow(log, "Не удалось сохранить AI-память")
    # Удалить теги из текста для отображения пользователю
    clean = _MEMORY_TAG_RE.sub("", text).rstrip()
    return clean, saved


async def _process_ai_turn(
    message: Message,
    state: FSMContext,
    pool: asyncpg.Pool,
    http: aiohttp.ClientSession,
    user_content: str,
    thinking_text: str = "🤖 <i>Разбираюсь...</i>",
) -> None:
    data = await state.get_data()
    messages: list = data.get("messages", [])
    turns: int = data.get("turns", 0)

    if turns >= _MAX_TURNS:
        await state.clear()
        await message.answer("Диалог стал слишком длинным. Начните новую сессию: /ai")
        return

    messages.append({"role": "user", "content": user_content})

    # ── Fast path: detect action intent without LLM ──────────────────────────
    fast_reply = await _fast_parse_action(user_content, pool, message.from_user.id)
    if fast_reply:
        thinking = await message.answer("⚡ <i>Готово!</i>", parse_mode="HTML")
        reply = fast_reply
    else:
        # Показываем индикатор набора текста перед вызовом API
        try:
            await message.bot.send_chat_action(message.chat.id, "typing")
        except Exception:
            pass

        thinking = await message.answer(thinking_text, parse_mode="HTML")

        try:
            reply = await asyncio.wait_for(
                _call_ai_providers(messages, pool, message.from_user.id, http),
                timeout=75.0,
            )
        except asyncio.TimeoutError:
            kb = InlineKeyboardBuilder()
            kb.button(text="🔄 Повторить", callback_data=AiCb(action="retry"))
            kb.button(text="❌ Выйти", callback_data=AiCb(action="stop"))
            kb.adjust(2)
            try:
                await thinking.edit_text(
                    "⏰ <b>AI-ассистент не ответил вовремя</b>\n\n"
                    "Все модели заняты или перегружены. Попробуйте ещё раз.",
                    parse_mode="HTML",
                    reply_markup=kb.as_markup(),
                )
            except Exception:
                await message.answer(
                    "⏰ <b>AI-ассистент не ответил вовремя.</b> Попробуйте ещё раз.",
                    parse_mode="HTML",
                    reply_markup=kb.as_markup(),
                )
            return

    if isinstance(reply, dict) and reply.get("__pending__"):
        # AI wants to perform an action — show confirmation
        action_data = reply["action_data"]
        ai_message = reply["ai_message"]
        preview = action_data.get("preview", "")

        await state.update_data(
            messages=messages + [{"role": "assistant", "content": ai_message}],
            turns=turns + 1,
            pending_action_data=action_data,
        )

        kb = InlineKeyboardBuilder()
        kb.button(text="✅ Подтвердить", callback_data=AiCb(action="confirm_action"))
        kb.button(text="❌ Отмена", callback_data=AiCb(action="cancel_action"))
        kb.adjust(2)

        # Build enhanced confirmation screen
        action_name = action_data.get("pending_action", "")
        preview = action_data.get("preview", "")

        icons = {
            "create_channel": "📡",
            "create_group": "👥",
            "create_bot": "🤖",
            "launch_broadcast": "📢",
            "post_to_channel": "📤",
            "update_bot_profile": "✏️",
            "schedule_broadcast": "⏱️",
            "bulk_create_channels": "📋",
        }
        icon = icons.get(action_name, "⚡")

        acc_info = ""
        if action_data.get("acc_id"):
            acc_row = await pool.fetchrow(
                "SELECT phone, first_name FROM tg_accounts WHERE id=$1",
                action_data["acc_id"],
            )
            if acc_row:
                acc_name = acc_row["first_name"] or acc_row["phone"]
                acc_info = f"\n📱 Аккаунт: <code>{acc_name}</code>"

        audience_info = ""
        if "audience" in action_data:
            audience_info = f"\n👥 Получателей: {action_data['audience']}"

        count_info = ""
        if action_data.get("count"):
            count_info = f"\n📊 Количество: {action_data['count']}"

        confirm_text = (
            f"{ai_message}\n\n"
            f"━━━━━━━━━━━━━━━━━━\n"
            f"{icon} <b>Подтвердите действие:</b>\n"
            f"<code>{preview}</code>"
            f"{acc_info}{audience_info}{count_info}"
        )
        await _edit_or_answer_long(
            thinking, message, confirm_text, reply_markup=kb.as_markup()
        )
    else:
        # Regular AI response (text)
        str_reply = reply if isinstance(reply, str) else str(reply)

        # Парсим и сохраняем [MEMORY: ...] теги перед отображением
        display_reply, saved_count = await _parse_and_save_memory(
            str_reply, pool, message.from_user.id
        )
        # Сохраняем в историю оригинал без тегов
        messages.append({"role": "assistant", "content": display_reply})
        await state.update_data(messages=messages, turns=turns + 1)

        if saved_count > 0:
            display_reply = (
                display_reply
                + f"\n\n💾 <i>Запомнил {saved_count} запись(-и) в памяти.</i>"
            )

        if "AI-ассистент временно недоступен" in display_reply:
            kb = InlineKeyboardBuilder()
            kb.button(text="🔄 Повторить", callback_data=AiCb(action="retry"))
            kb.button(text="🏠 Меню", callback_data=BmCb(action="main"))
            kb.adjust(2)
            await _edit_or_answer_long(
                thinking, message, display_reply, reply_markup=kb.as_markup()
            )
            return

        kb = InlineKeyboardBuilder()
        kb.button(text="📚 Память", callback_data=AiCb(action="memory"))
        kb.button(text="🗑 Очистить историю", callback_data=AiCb(action="clear_history"))
        kb.button(text="❌ Выйти из чата", callback_data=AiCb(action="stop"))
        kb.adjust(3)
        await _edit_or_answer_long(
            thinking, message, display_reply, reply_markup=kb.as_markup()
        )


@router.message(Command("ai"))
async def cmd_ai(message: Message) -> None:
    from bot.callbacks import BmCb, AiCb

    kb = InlineKeyboardBuilder()
    kb.button(text="🤖 Открыть ИИ Помощник", callback_data=AiCb(action="start"))
    kb.button(text="🏠 Главное меню", callback_data=BmCb(action="main"))
    kb.adjust(1)
    await message.answer(
        "🤖 <b>ИИ Помощник</b>\n\n"
        "Нейросетевой ассистент для создания контента и управления ботами.\n\n"
        "Находится в: <b>Настройки → 🤖 ИИ Помощник</b>",
        reply_markup=kb.as_markup(),
        parse_mode="HTML",
    )


@router.message(Command("remember"))
async def cmd_remember(message: Message, pool: asyncpg.Pool) -> None:
    text = message.text or ""
    body = text.split(maxsplit=1)[1].strip() if len(text.split(maxsplit=1)) > 1 else ""
    if not body:
        await message.answer(
            "Напиши так: <code>/remember что важно запомнить #тег</code>",
            parse_mode="HTML",
        )
        return

    tags = [part[1:] for part in body.split() if part.startswith("#") and len(part) > 1]
    title = body[:80].strip()
    try:
        item = await ai_memory.remember(
            pool,
            message.from_user.id,
            body,
            title=title,
            tags=tags,
            source="command",
        )
    except Exception:
        log.exception("Failed to save AI memory")
        await message.answer("❌ Не удалось сохранить запись. Проверьте логи.")
        return

    await message.answer(
        f"✅ Запомнил. ID <code>#{item.id}</code>",
        parse_mode="HTML",
    )


@router.message(Command("memory"))
async def cmd_memory(message: Message, pool: asyncpg.Pool) -> None:
    text = message.text or ""
    query = text.split(maxsplit=1)[1].strip() if len(text.split(maxsplit=1)) > 1 else ""
    try:
        items = await ai_memory.search(pool, message.from_user.id, query, limit=10)
    except Exception:
        log.exception("Failed to read AI memory")
        await message.answer("❌ Не удалось прочитать память. Проверьте логи.")
        return

    await message.answer(ai_memory.format_for_user(items), parse_mode="HTML")


@router.message(Command("forget"))
async def cmd_forget(message: Message, pool: asyncpg.Pool) -> None:
    text = message.text or ""
    raw_id = (
        text.split(maxsplit=1)[1].strip().lstrip("#")
        if len(text.split(maxsplit=1)) > 1
        else ""
    )
    if not raw_id.isdigit():
        await message.answer("Напиши так: <code>/forget 123</code>", parse_mode="HTML")
        return

    try:
        deleted = await ai_memory.delete(pool, message.from_user.id, int(raw_id))
    except Exception:
        log.exception("Failed to delete AI memory")
        await message.answer("❌ Не удалось удалить запись памяти. Проверьте логи.")
        return

    await message.answer(
        "✅ Удалил." if deleted else "Такой записи в твоей памяти нет."
    )


@router.callback_query(AiCb.filter(F.action == "start"))
async def cb_ai_start(
    callback: CallbackQuery, state: FSMContext, pool: asyncpg.Pool
) -> None:
    await safe_answer(callback)
    if not await require_plan(pool, callback.from_user.id, "enterprise"):
        await callback.message.edit_text(
            "🔒 <b>AI-ассистент — 💎 ПОДПИСКА</b>\n\nОформите подписку: /subscription",
            parse_mode="HTML",
        )
        return
    await state.set_state(AiChat.chatting)
    await state.update_data(messages=[], turns=0)
    kb = InlineKeyboardBuilder()
    kb.button(text="📚 Память", callback_data=AiCb(action="memory"))
    kb.button(text="🗑 Очистить историю", callback_data=AiCb(action="clear_history"))
    kb.button(text="❌ Выйти из чата", callback_data=AiCb(action="stop"))
    kb.adjust(3)
    await callback.message.edit_text(
        "🤖 <b>AI-ассистент запущен</b>\n\n"
        "Задайте вопрос или поставьте задачу. Я могу анализировать данные ваших ботов, "
        "читать файлы, запускать рассылки, обновлять профили и публиковать посты в каналы.\n\n"
        "⚠️ Любое действие потребует вашего подтверждения.",
        parse_mode="HTML",
        reply_markup=kb.as_markup(),
    )


@router.callback_query(AiCb.filter(F.action == "stop"))
async def cb_ai_stop(callback: CallbackQuery, state: FSMContext) -> None:
    await safe_answer(callback)
    await state.clear()
    kb = InlineKeyboardBuilder()
    kb.button(text="🤖 Новая сессия", callback_data=AiCb(action="start"))
    kb.button(text="🏠 Главное меню", callback_data=BmCb(action="main"))
    kb.adjust(1)
    await callback.message.edit_text(
        "✅ Сессия AI-ассистента завершена.",
        reply_markup=kb.as_markup(),
    )


@router.callback_query(AiCb.filter(F.action == "clear_history"))
async def cb_ai_clear_history(callback: CallbackQuery, state: FSMContext) -> None:
    await callback.answer("🗑 История очищена", show_alert=False)
    # Сбрасываем историю диалога, но остаёмся в режиме чата
    await state.set_state(AiChat.chatting)
    await state.update_data(messages=[], turns=0)
    kb = InlineKeyboardBuilder()
    kb.button(text="🗑 Очистить историю", callback_data=AiCb(action="clear_history"))
    kb.button(text="❌ Выйти из чата", callback_data=AiCb(action="stop"))
    kb.adjust(2)
    await callback.message.edit_text(
        "🗑 <b>История диалога очищена</b>\n\n"
        "Контекст сброшен — можете начать новый диалог с чистого листа.",
        parse_mode="HTML",
        reply_markup=kb.as_markup(),
    )


@router.callback_query(AiCb.filter(F.action == "confirm_action"))
async def cb_ai_confirm_action(
    callback: CallbackQuery,
    state: FSMContext,
    pool: asyncpg.Pool,
    http: aiohttp.ClientSession,
) -> None:
    await safe_answer(callback)
    data = await state.get_data()
    action_data = data.get("pending_action_data")
    if not action_data:
        await callback.message.edit_text(
            "⚠️ Данные действия устарели. Попросите ассистента повторить."
        )
        return

    status_msg = await callback.message.edit_text(
        "⏳ <i>Выполняю...</i>", parse_mode="HTML"
    )

    async def _do_action() -> None:
        try:
            result = await execute_action(
                action_data, pool, callback.from_user.id, http
            )
        except Exception as exc:
            result = f"❌ Ошибка при выполнении: {escape(str(exc)[:200])}"
        await state.update_data(pending_action_data=None)
        kb = InlineKeyboardBuilder()
        kb.button(text="🗑 Очистить историю", callback_data=AiCb(action="clear_history"))
        kb.button(text="❌ Выйти из чата", callback_data=AiCb(action="stop"))
        kb.adjust(2)
        try:
            await status_msg.edit_text(
                f"<b>Результат выполнения:</b>\n\n{result}",
                parse_mode="HTML",
                reply_markup=kb.as_markup(),
            )
        except Exception:
            try:
                await callback.message.answer(
                    f"<b>Результат выполнения:</b>\n\n{result}",
                    parse_mode="HTML",
                    reply_markup=kb.as_markup(),
                )
            except Exception:
                pass

    asyncio.create_task(_do_action())


@router.callback_query(AiCb.filter(F.action == "cancel_action"))
async def cb_ai_cancel_action(callback: CallbackQuery, state: FSMContext) -> None:
    await callback.answer("Действие отменено")
    await state.update_data(pending_action_data=None)
    kb = InlineKeyboardBuilder()
    kb.button(text="🗑 Очистить историю", callback_data=AiCb(action="clear_history"))
    kb.button(text="❌ Выйти из чата", callback_data=AiCb(action="stop"))
    kb.adjust(2)
    await callback.message.edit_text(
        "❌ Действие отменено. Можете задать новый вопрос или задачу.",
        reply_markup=kb.as_markup(),
    )


@router.callback_query(AiCb.filter(F.action == "memory"))
async def cb_ai_memory(
    callback: CallbackQuery,
    pool: asyncpg.Pool,
) -> None:
    """Показать список последних 10 записей памяти пользователя."""
    await safe_answer(callback)
    try:
        items = await ai_memory.search(pool, callback.from_user.id, "", limit=10)
    except Exception:
        log.exception("Failed to load AI memory list")
        await callback.message.edit_text(
            "❌ Не удалось загрузить память.",
            reply_markup=InlineKeyboardBuilder()
            .button(text="◀️ Назад", callback_data=AiCb(action="start"))
            .as_markup(),
        )
        return

    kb = InlineKeyboardBuilder()
    if items:
        lines = ["📚 <b>Память AI-ассистента</b> (последние 10)\n"]
        for item in items:
            pin = "📌 " if item.pinned else ""
            kind_label = {"note": "📝", "insight": "💡", "reminder": "⏰"}.get(
                item.kind, "📄"
            )
            title_text = (
                escape(item.title[:50]) if item.title else escape(item.body[:50])
            )
            dt = item.created_at.strftime("%d.%m") if item.created_at else ""
            lines.append(f"{kind_label} {pin}<b>{title_text}</b> <i>{dt}</i>")
            kb.button(
                text=f"🗑 #{item.id} {(item.title or item.body[:30])[:30]}",
                callback_data=AiCb(action="memory_delete", memory_id=item.id),
            )
        kb.adjust(1)
        text = "\n".join(lines)
    else:
        text = "📚 <b>Память пустая</b>\n\nAI-ассистент будет сохранять важные сведения о вас здесь."

    kb.button(text="◀️ Назад в чат", callback_data=AiCb(action="start"))
    await callback.message.edit_text(
        text, parse_mode="HTML", reply_markup=kb.as_markup()
    )


@router.callback_query(AiCb.filter(F.action == "memory_delete"))
async def cb_ai_memory_delete(
    callback: CallbackQuery,
    callback_data: AiCb,
    pool: asyncpg.Pool,
) -> None:
    """Удалить запись из памяти AI."""
    await safe_answer(callback)
    memory_id = callback_data.memory_id
    if not memory_id:
        await callback.answer("❌ Некорректный ID записи", show_alert=True)
        return
    try:
        deleted = await ai_memory.delete(pool, callback.from_user.id, memory_id)
    except Exception:
        log.exception("Failed to delete AI memory item %s", memory_id)
        await callback.answer("❌ Ошибка при удалении", show_alert=True)
        return

    if deleted:
        await callback.answer("🗑 Запись удалена", show_alert=False)
    else:
        await callback.answer("Запись не найдена", show_alert=False)

    # Перезагрузить список памяти
    try:
        items = await ai_memory.search(pool, callback.from_user.id, "", limit=10)
    except Exception:
        items = []

    kb = InlineKeyboardBuilder()
    if items:
        lines = ["📚 <b>Память AI-ассистента</b> (последние 10)\n"]
        for item in items:
            pin = "📌 " if item.pinned else ""
            kind_label = {"note": "📝", "insight": "💡", "reminder": "⏰"}.get(
                item.kind, "📄"
            )
            title_text = (
                escape(item.title[:50]) if item.title else escape(item.body[:50])
            )
            dt = item.created_at.strftime("%d.%m") if item.created_at else ""
            lines.append(f"{kind_label} {pin}<b>{title_text}</b> <i>{dt}</i>")
            kb.button(
                text=f"🗑 #{item.id} {(item.title or item.body[:30])[:30]}",
                callback_data=AiCb(action="memory_delete", memory_id=item.id),
            )
        kb.adjust(1)
        text = "\n".join(lines)
    else:
        text = "📚 <b>Память пустая</b>\n\nAI-ассистент будет сохранять важные сведения о вас здесь."

    kb.button(text="◀️ Назад в чат", callback_data=AiCb(action="start"))
    await callback.message.edit_text(
        text, parse_mode="HTML", reply_markup=kb.as_markup()
    )


@router.callback_query(AiCb.filter(F.action == "retry"))
async def cb_ai_retry(
    callback: CallbackQuery,
    state: FSMContext,
    pool: asyncpg.Pool,
    http: aiohttp.ClientSession,
) -> None:
    await safe_answer(callback)
    data = await state.get_data()
    messages: list = data.get("messages", [])
    if not messages:
        await callback.message.edit_text("❌ Нет предыдущего сообщения для повтора.")
        return
    # Find last user message
    last_user = next((m for m in reversed(messages) if m.get("role") == "user"), None)
    if not last_user:
        await callback.message.edit_text("❌ Нет предыдущего сообщения для повтора.")
        return
    # Remove last assistant message if present so we don't duplicate it
    retry_messages = (
        messages[:-1] if messages[-1].get("role") == "assistant" else messages
    )
    await callback.message.edit_text("⏳ <b>Повторяю запрос...</b>", parse_mode="HTML")
    try:
        reply = await asyncio.wait_for(
            _call_ai_providers(retry_messages, pool, callback.from_user.id, http),
            timeout=75.0,
        )
    except asyncio.TimeoutError:
        kb_to = InlineKeyboardBuilder()
        kb_to.button(text="🔄 Повторить", callback_data=AiCb(action="retry"))
        kb_to.button(text="❌ Выйти", callback_data=AiCb(action="stop"))
        kb_to.adjust(2)
        await callback.message.edit_text(
            "⏰ <b>AI-ассистент не ответил вовремя.</b> Попробуйте ещё раз.",
            parse_mode="HTML",
            reply_markup=kb_to.as_markup(),
        )
        return
    if isinstance(reply, str):
        str_reply = reply
    else:
        str_reply = str(reply)
    kb = InlineKeyboardBuilder()
    if "AI-ассистент временно недоступен" in str_reply:
        kb.button(text="🔄 Повторить", callback_data=AiCb(action="retry"))
        kb.button(text="🏠 Меню", callback_data=BmCb(action="main"))
        kb.adjust(2)
    else:
        kb.button(text="🗑 Очистить историю", callback_data=AiCb(action="clear_history"))
        kb.button(text="❌ Выйти из чата", callback_data=AiCb(action="stop"))
        kb.adjust(2)
    await _edit_or_answer_long(
        callback.message, callback.message, str_reply, reply_markup=kb.as_markup()
    )


@router.message(AiChat.chatting, F.text)
async def msg_ai_chat(
    message: Message,
    state: FSMContext,
    pool: asyncpg.Pool,
    http: aiohttp.ClientSession,
) -> None:
    await _process_ai_turn(message, state, pool, http, message.text or "")


@router.message(AiChat.chatting, F.document)
async def msg_ai_document(
    message: Message,
    state: FSMContext,
    pool: asyncpg.Pool,
    http: aiohttp.ClientSession,
) -> None:
    try:
        file_name, file_text = await _read_attached_document(message)
    except ValueError as exc:
        await message.answer(
            "📎 <b>Не могу прочитать этот файл</b>\n\n"
            f"{escape(str(exc))}\n\n"
            f"Можно отправить: <code>{_SUPPORTED_FILE_HINT}</code>",
            parse_mode="HTML",
        )
        return
    except Exception:
        log.exception("Failed to read AI document")
        await message.answer(
            "📎 <b>Не получилось прочитать файл</b>\n\n"
            "Попробуйте отправить его еще раз или сохраните текст в формате TXT/MD.",
            parse_mode="HTML",
        )
        return

    prompt = _build_file_prompt(file_name, file_text, message.caption)
    safe_name = escape(file_name)
    await _process_ai_turn(
        message,
        state,
        pool,
        http,
        prompt,
        thinking_text=f"📎 <i>Читаю файл {safe_name}...</i>",
    )
