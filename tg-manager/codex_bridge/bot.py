from __future__ import annotations

import asyncio
import html
import json
import logging
import os
import shutil
import tempfile
import time
import uuid
from io import BytesIO
from pathlib import Path

from aiogram import Bot, Dispatcher, F
from aiogram.client.default import DefaultBotProperties
from aiogram.enums import ChatAction, ParseMode
from aiogram.filters import Command
from aiogram.types import CallbackQuery, InlineKeyboardButton, InlineKeyboardMarkup, Message
from dotenv import load_dotenv

BASE_DIR = Path(__file__).resolve().parent
STATE_DIR = BASE_DIR / ".state"
HISTORY_DIR = STATE_DIR / "history"
LOG = logging.getLogger("codex_bridge")

TEXT_FILE_EXTENSIONS = {
    ".cfg", ".conf", ".css", ".csv", ".env", ".html", ".ini", ".js", ".json",
    ".jsx", ".log", ".md", ".mjs", ".prisma", ".py", ".sh", ".sql", ".toml",
    ".ts", ".tsx", ".txt", ".xml", ".yaml", ".yml",
}
STRUCTURED_FILE_EXTENSIONS = {".pdf", ".docx", ".xlsx"}


def _str_env(name: str, default: str = "") -> str:
    return os.getenv(name, default).strip()


def _bool_env(name: str, default: bool = False) -> bool:
    raw = os.getenv(name)
    if raw is None:
        return default
    value = raw.strip().lower()
    if value in {"1", "true", "yes", "on"}:
        return True
    if value in {"0", "false", "no", "off", ""}:
        return False
    raise RuntimeError(f"Invalid boolean {name}={raw!r}; use true/false")


def _int_env(name: str, default: int) -> int:
    raw = os.getenv(name)
    if raw is None or not raw.strip():
        return default
    try:
        return int(raw)
    except ValueError as exc:
        raise RuntimeError(f"Invalid integer {name}={raw!r}") from exc


def _ids_env(name: str) -> set[int]:
    ids: set[int] = set()
    for item in os.getenv(name, "").split(","):
        item = item.strip()
        if not item:
            continue
        try:
            ids.add(int(item))
        except ValueError as exc:
            raise RuntimeError(f"Invalid Telegram user id in {name}: {item!r}") from exc
    return ids


def _resolve_project_dir() -> Path:
    raw = _str_env("CODEX_PROJECT_DIR", "..")
    path = Path(raw)
    if not path.is_absolute():
        path = (BASE_DIR / path).resolve()
    return path


def _resolve_codex_cmd() -> str:
    configured = _str_env("CODEX_CMD")
    if configured:
        return configured

    appdata = os.getenv("APPDATA")
    if appdata:
        npm_codex = Path(appdata) / "npm" / "codex.cmd"
        if npm_codex.exists():
            return str(npm_codex)

    discovered = shutil.which("codex.cmd") or shutil.which("codex")
    if discovered and "WindowsApps" not in discovered:
        return discovered

    raise RuntimeError("Codex CLI was not found. Set CODEX_CMD in .env.")


def _history_path(user_id: int) -> Path:
    return HISTORY_DIR / f"{user_id}.json"


def _load_history(user_id: int, turns: int) -> list[dict[str, str]]:
    path = _history_path(user_id)
    if not path.exists():
        return []
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return []
    if not isinstance(data, list):
        return []
    cleaned = [
        item
        for item in data
        if isinstance(item, dict)
        and isinstance(item.get("role"), str)
        and isinstance(item.get("content"), str)
    ]
    return cleaned[-turns * 2 :]


def _save_history(user_id: int, history: list[dict[str, str]], turns: int) -> None:
    HISTORY_DIR.mkdir(parents=True, exist_ok=True)
    _history_path(user_id).write_text(
        json.dumps(history[-turns * 2 :], ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def _clear_history(user_id: int) -> None:
    path = _history_path(user_id)
    if path.exists():
        path.unlink()


def _split_telegram_text(text: str, limit: int = 3900) -> list[str]:
    if len(text) <= limit:
        return [text]

    parts: list[str] = []
    remaining = text
    while remaining:
        chunk = remaining[:limit]
        split_at = max(chunk.rfind("\n\n"), chunk.rfind("\n"), chunk.rfind(" "))
        if split_at < limit // 2:
            split_at = limit
        parts.append(remaining[:split_at].strip())
        remaining = remaining[split_at:].strip()
    return [part for part in parts if part]


def _decode_text_bytes(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "utf-16", "cp1251", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _extract_pdf_text(raw: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF пока не читается: не установлена библиотека pypdf.") from exc

    reader = PdfReader(BytesIO(raw))
    parts: list[str] = []
    max_pages = 20
    for index in range(min(len(reader.pages), max_pages)):
        page_text = reader.pages[index].extract_text() or ""
        if page_text.strip():
            parts.append(f"[Страница {index + 1}]\n{page_text.strip()}")
    if len(reader.pages) > max_pages:
        parts.append(f"[Показаны первые {max_pages} страниц из {len(reader.pages)}.]")
    return "\n\n".join(parts)


def _extract_docx_text(raw: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX пока не читается: не установлена библиотека python-docx.") from exc

    doc = Document(BytesIO(raw))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table_index, table in enumerate(doc.tables, start=1):
        rows: list[str] = []
        for row in table.rows[:100]:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append(f"[Таблица {table_index}]\n" + "\n".join(rows))
    return "\n\n".join(parts)


def _extract_xlsx_text(raw: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError("XLSX пока не читается: не установлена библиотека openpyxl.") from exc

    workbook = load_workbook(BytesIO(raw), read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for sheet in workbook.worksheets[:5]:
            rows: list[str] = []
            for row in sheet.iter_rows(max_row=120, max_col=20, values_only=True):
                values = ["" if value is None else str(value) for value in row]
                if any(value.strip() for value in values):
                    rows.append("\t".join(values).rstrip())
            if rows:
                parts.append(f"[Лист: {sheet.title}]\n" + "\n".join(rows))
        if len(workbook.worksheets) > 5:
            parts.append(f"[Показаны первые 5 листов из {len(workbook.worksheets)}.]")
    finally:
        workbook.close()
    return "\n\n".join(parts)


def _extract_file_text(file_name: str, mime_type: str | None, raw: bytes) -> str:
    suffix = Path(file_name).suffix.lower()
    mime_type = (mime_type or "").lower()
    if suffix in TEXT_FILE_EXTENSIONS or mime_type.startswith("text/"):
        return _decode_text_bytes(raw)
    if suffix == ".pdf":
        return _extract_pdf_text(raw)
    if suffix == ".docx":
        return _extract_docx_text(raw)
    if suffix == ".xlsx":
        return _extract_xlsx_text(raw)
    allowed = sorted(TEXT_FILE_EXTENSIONS | STRUCTURED_FILE_EXTENSIONS)
    raise RuntimeError("Пока я читаю такие файлы: " + ", ".join(allowed))


class BridgeConfig:
    def __init__(self) -> None:
        load_dotenv(BASE_DIR / ".env")

        self.bot_token = _str_env("TELEGRAM_CODEX_BOT_TOKEN")
        if not self.bot_token:
            raise RuntimeError("Missing TELEGRAM_CODEX_BOT_TOKEN in tg-manager/codex_bridge/.env")

        self.admin_user_ids = _ids_env("TELEGRAM_ADMIN_USER_IDS")
        self.admin_user_ids.update(_ids_env("TELEGRAM_OWNER_USER_ID"))
        self.allowed_user_ids = _ids_env("TELEGRAM_ALLOWED_USER_IDS")
        self.allow_all_users = _bool_env("ALLOW_ALL_TELEGRAM_USERS", False)
        self.project_dir = _resolve_project_dir()
        self.codex_cmd = _resolve_codex_cmd()
        self.model = _str_env("CODEX_MODEL")
        self.sandbox = _str_env("CODEX_SANDBOX", "workspace-write")
        self.approval = _str_env("CODEX_APPROVAL", "never")
        self.timeout_seconds = _int_env("CODEX_TIMEOUT_SECONDS", 900)
        self.history_turns = _int_env("CODEX_HISTORY_TURNS", 8)
        self.confirm_before_run = _bool_env("CODEX_CONFIRM_BEFORE_RUN", False)
        self.max_file_bytes = _int_env("TELEGRAM_MAX_FILE_BYTES", 1_048_576)
        self.max_file_chars = _int_env("TELEGRAM_MAX_FILE_CHARS", 60_000)
        self.progress_seconds = _int_env("TELEGRAM_PROGRESS_SECONDS", 300)

        if self.progress_seconds < 60:
            self.progress_seconds = 60
        if not self.project_dir.exists():
            raise RuntimeError(f"CODEX_PROJECT_DIR does not exist: {self.project_dir}")

    def is_allowed(self, user_id: int) -> bool:
        return (
            self.allow_all_users
            or user_id in self.admin_user_ids
            or user_id in self.allowed_user_ids
        )

    def role_for(self, user_id: int) -> str:
        if user_id in self.admin_user_ids:
            return "админ"
        if self.is_allowed(user_id):
            return "доступ есть"
        return "нет доступа"


class CodexBridge:
    def __init__(self, config: BridgeConfig) -> None:
        self.config = config
        self.lock = asyncio.Lock()
        self.current_process: asyncio.subprocess.Process | None = None
        self.current_started_at: float | None = None
        self.current_user_id: int | None = None
        self.current_text: str | None = None

    def is_busy(self) -> bool:
        return self.lock.locked() or self.current_process is not None

    def build_prompt(self, user_id: int, text: str) -> str:
        history = _load_history(user_id, self.config.history_turns)
        history_text = "\n".join(
            f"{item['role'].upper()}: {item['content']}" for item in history
        )
        history_block = f"\nRecent Telegram conversation:\n{history_text}\n" if history_text else ""

        return (
            "You are Codex working through the BotMother Telegram bridge.\n"
            "All user-facing messages must be in Russian, clear, direct, and human-sounding.\n"
            "Do not use repetitive template phrases like 'Принял. Начинаю работать над проектом'.\n"
            "When changing code, inspect files first, edit carefully, and verify when possible.\n"
            "Only edit files inside the current working directory unless the user explicitly says otherwise.\n"
            "Do not ask for interactive approvals; work within the configured sandbox.\n"
            f"Working directory: {self.config.project_dir}\n"
            f"{history_block}\n"
            f"Telegram message:\n{text}\n"
        )

    async def run_codex(self, prompt: str) -> str:
        with tempfile.NamedTemporaryFile(
            mode="w", encoding="utf-8", suffix=".txt", delete=False
        ) as temp:
            output_path = Path(temp.name)

        command = [
            self.config.codex_cmd,
            "--ask-for-approval",
            self.config.approval,
            "exec",
            "--cd",
            str(self.config.project_dir),
            "--sandbox",
            self.config.sandbox,
            "--output-last-message",
            str(output_path),
        ]
        if self.config.model:
            command.extend(["--model", self.config.model])
        command.append("-")

        process = await asyncio.create_subprocess_exec(
            *command,
            stdin=asyncio.subprocess.PIPE,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
            cwd=str(self.config.project_dir),
        )
        self.current_process = process
        self.current_started_at = time.time()
        try:
            stdout, stderr = await asyncio.wait_for(
                process.communicate(prompt.encode("utf-8")),
                timeout=self.config.timeout_seconds,
            )
        except asyncio.TimeoutError:
            process.kill()
            await process.communicate()
            raise RuntimeError("Время ожидания вышло. Я остановил задачу, чтобы не зависнуть.")
        finally:
            self.current_process = None
            self.current_started_at = None
            self.current_user_id = None
            self.current_text = None

        if process.returncode != 0:
            detail = stderr.decode("utf-8", errors="replace").strip()
            if not detail:
                detail = stdout.decode("utf-8", errors="replace").strip()
            raise RuntimeError(detail or f"Codex завершился с ошибкой {process.returncode}")

        try:
            answer = output_path.read_text(encoding="utf-8").strip()
        finally:
            try:
                output_path.unlink()
            except OSError:
                pass

        if not answer:
            answer = stdout.decode("utf-8", errors="replace").strip()
        return answer.strip() or "Готово. Работа завершена."

    async def ask(self, user_id: int, text: str) -> str:
        async with self.lock:
            self.current_user_id = user_id
            self.current_text = text
            answer = await self.run_codex(self.build_prompt(user_id, text))
            history = _load_history(user_id, self.config.history_turns)
            history.extend(
                [
                    {"role": "user", "content": text},
                    {"role": "assistant", "content": answer},
                ]
            )
            _save_history(user_id, history, self.config.history_turns)
            return answer

    def busy_status(self) -> str:
        if self.current_process is None or self.current_started_at is None:
            return "Сейчас активной задачи нет."

        elapsed = int(time.time() - self.current_started_at)
        minutes = elapsed // 60
        seconds = elapsed % 60
        preview = (self.current_text or "").strip().replace("\n", " ")
        if len(preview) > 180:
            preview = preview[:177] + "..."
        return (
            f"Я на задаче уже {minutes} мин {seconds} сек.\n"
            f"Задача: {preview or 'без текста'}"
        )

    async def stop_current(self) -> bool:
        process = self.current_process
        if process is None or process.returncode is not None:
            return False
        process.kill()
        await process.wait()
        return True


class PendingRun:
    def __init__(self, user_id: int, chat_id: int, text: str) -> None:
        self.user_id = user_id
        self.chat_id = chat_id
        self.text = text
        self.created_at = time.time()


class ConfirmationStore:
    def __init__(self, ttl_seconds: int = 900) -> None:
        self.ttl_seconds = ttl_seconds
        self.pending: dict[str, PendingRun] = {}

    def create(self, user_id: int, chat_id: int, text: str) -> str:
        self.cleanup()
        run_id = uuid.uuid4().hex
        self.pending[run_id] = PendingRun(user_id, chat_id, text)
        return run_id

    def pop(self, run_id: str, user_id: int) -> PendingRun | None:
        self.cleanup()
        run = self.pending.get(run_id)
        if run is None or run.user_id != user_id:
            return None
        return self.pending.pop(run_id)

    def cancel(self, run_id: str, user_id: int) -> bool:
        self.cleanup()
        run = self.pending.get(run_id)
        if run is None or run.user_id != user_id:
            return False
        self.pending.pop(run_id)
        return True

    def cleanup(self) -> None:
        cutoff = time.time() - self.ttl_seconds
        for run_id, run in list(self.pending.items()):
            if run.created_at < cutoff:
                self.pending.pop(run_id, None)


def _confirm_markup(run_id: str) -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[
            [InlineKeyboardButton(text="Запустить", callback_data=f"codex:run:{run_id}")],
            [InlineKeyboardButton(text="Отменить", callback_data=f"codex:cancel:{run_id}")],
        ]
    )


async def send_long_message(message: Message, text: str) -> None:
    for part in _split_telegram_text(text):
        await message.answer(part, parse_mode=None)


async def typing_loop(bot: Bot, chat_id: int, stop: asyncio.Event) -> None:
    while not stop.is_set():
        await bot.send_chat_action(chat_id, ChatAction.TYPING)
        try:
            await asyncio.wait_for(stop.wait(), timeout=4)
        except asyncio.TimeoutError:
            continue


async def progress_loop(bot: Bot, chat_id: int, stop: asyncio.Event, interval: int) -> None:
    phrases = [
        "Я в работе. Не завис, просто задача длинная.",
        "Двигаюсь дальше. Без лишнего шума, но держу курс.",
        "Работа продолжается. Следующий статус дам, если снова затянется.",
    ]
    try:
        await asyncio.wait_for(stop.wait(), timeout=interval)
        return
    except asyncio.TimeoutError:
        pass

    elapsed = interval
    index = 0
    while not stop.is_set():
        minutes = max(1, elapsed // 60)
        await bot.send_message(chat_id, f"{phrases[index % len(phrases)]}\nПрошло примерно {minutes} мин.")
        index += 1
        try:
            await asyncio.wait_for(stop.wait(), timeout=interval)
        except asyncio.TimeoutError:
            elapsed += interval


async def run_user_request(
    bot: Bot,
    bridge: CodexBridge,
    message: Message,
    user_id: int,
    text: str,
) -> None:
    if bridge.is_busy():
        await message.answer(
            "Я уже выполняю задачу. Вторую параллельно не запускаю, чтобы не смешать изменения.\n\n"
            "Проверить ход: /busy\n"
            "Остановить текущую: /stop"
        )
        return

    await message.answer("Ок, беру в работу. Если задача затянется, дам короткий статус примерно раз в 5 минут.")
    stop_event = asyncio.Event()
    typing_task = asyncio.create_task(typing_loop(bot, message.chat.id, stop_event))
    progress_task = asyncio.create_task(
        progress_loop(bot, message.chat.id, stop_event, bridge.config.progress_seconds)
    )
    try:
        answer = await bridge.ask(user_id, text)
        await send_long_message(message, answer)
    except Exception as exc:
        LOG.exception("Codex request failed")
        await message.answer(
            "Не получилось выполнить задачу:\n"
            f"<code>{html.escape(str(exc)[:3500])}</code>"
        )
    finally:
        stop_event.set()
        await typing_task
        await progress_task


async def read_document_text(bot: Bot, message: Message, config: BridgeConfig) -> tuple[str, str]:
    document = message.document
    if document is None:
        raise RuntimeError("Файл не найден в сообщении.")

    file_name = document.file_name or "telegram-file.txt"
    file_size = document.file_size or 0
    if file_size > config.max_file_bytes:
        limit_mb = config.max_file_bytes / 1024 / 1024
        raise RuntimeError(f"Файл слишком большой. Сейчас лимит {limit_mb:.1f} МБ.")

    telegram_file = await bot.get_file(document.file_id)
    if telegram_file.file_path is None:
        raise RuntimeError("Telegram не отдал путь к файлу.")

    buffer = BytesIO()
    await bot.download_file(telegram_file.file_path, destination=buffer)
    raw = buffer.getvalue()
    if len(raw) > config.max_file_bytes:
        limit_mb = config.max_file_bytes / 1024 / 1024
        raise RuntimeError(f"Файл слишком большой. Сейчас лимит {limit_mb:.1f} МБ.")

    text = _extract_file_text(file_name, document.mime_type, raw).replace("\x00", "").strip()
    if not text:
        raise RuntimeError("В файле не нашёл читаемый текст.")
    if len(text) > config.max_file_chars:
        text = (
            text[: config.max_file_chars].rstrip()
            + "\n\n[Файл длинный, дальше текст обрезан. Пришлите нужную часть отдельно, если нужны детали ниже.]"
        )
    return file_name, text


def _access_denied_text(user_id: int) -> str:
    return (
        "У этого Telegram-аккаунта нет доступа к боту.\n\n"
        f"Telegram ID: <code>{user_id}</code>\n\n"
        "Добавьте этот ID в разрешённые пользователи и перезапустите мост."
    )


async def main() -> None:
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(name)s] %(levelname)s: %(message)s",
    )
    config = BridgeConfig()
    bridge = CodexBridge(config)
    confirmations = ConfirmationStore(ttl_seconds=config.timeout_seconds)

    bot = Bot(
        token=config.bot_token,
        default=DefaultBotProperties(parse_mode=ParseMode.HTML),
    )
    dp = Dispatcher()

    @dp.message(Command("start"))
    async def start(message: Message) -> None:
        user_id = message.from_user.id
        if not config.is_allowed(user_id):
            await message.answer(_access_denied_text(user_id))
            return

        await message.answer(
            "Я на месте.\n\n"
            f"Telegram ID: <code>{user_id}</code>\n"
            f"Роль: <code>{config.role_for(user_id)}</code>\n"
            f"Рабочая папка: <code>{html.escape(str(config.project_dir))}</code>\n\n"
            "Пиши задачу обычным сообщением или отправляй файл."
        )

    @dp.message(Command("status"))
    async def status(message: Message) -> None:
        user_id = message.from_user.id
        if not config.is_allowed(user_id):
            await message.answer(_access_denied_text(user_id))
            return

        await message.answer(
            "Мост работает.\n"
            f"Роль: <code>{config.role_for(user_id)}</code>\n"
            f"Рабочая папка: <code>{html.escape(str(config.project_dir))}</code>\n"
            f"Codex: <code>{html.escape(config.codex_cmd)}</code>\n"
            f"Статусы: раз в <code>{config.progress_seconds}</code> сек.\n"
            f"Подтверждение перед запуском: <code>{config.confirm_before_run}</code>"
        )

    @dp.message(Command("reset"))
    async def reset(message: Message) -> None:
        user_id = message.from_user.id
        if not config.is_allowed(user_id):
            await message.answer(_access_denied_text(user_id))
            return

        _clear_history(user_id)
        await message.answer("Память диалога очищена.")

    @dp.message(Command("busy"))
    async def busy(message: Message) -> None:
        user_id = message.from_user.id
        if not config.is_allowed(user_id):
            await message.answer(_access_denied_text(user_id))
            return

        await message.answer(html.escape(bridge.busy_status()))

    @dp.message(Command("stop"))
    async def stop(message: Message) -> None:
        user_id = message.from_user.id
        if not config.is_allowed(user_id):
            await message.answer(_access_denied_text(user_id))
            return

        stopped = await bridge.stop_current()
        await message.answer("Остановил текущую задачу." if stopped else "Сейчас активной задачи нет.")

    @dp.callback_query(F.data.startswith("codex:cancel:"))
    async def cancel_run(callback: CallbackQuery) -> None:
        user_id = callback.from_user.id
        if not config.is_allowed(user_id):
            await callback.answer("Нет доступа", show_alert=True)
            return

        run_id = callback.data.rsplit(":", 1)[-1]
        if confirmations.cancel(run_id, user_id):
            await callback.answer("Отменено")
            if callback.message:
                await callback.message.edit_text("Запуск отменён.")
        else:
            await callback.answer("Задача уже не актуальна", show_alert=True)

    @dp.callback_query(F.data.startswith("codex:run:"))
    async def confirm_run(callback: CallbackQuery) -> None:
        user_id = callback.from_user.id
        if not config.is_allowed(user_id):
            await callback.answer("Нет доступа", show_alert=True)
            return

        run_id = callback.data.rsplit(":", 1)[-1]
        run = confirmations.pop(run_id, user_id)
        if run is None:
            await callback.answer("Задача уже не актуальна", show_alert=True)
            return
        if bridge.is_busy():
            await callback.answer("Я уже занят другой задачей", show_alert=True)
            return

        await callback.answer("Запускаю")
        if callback.message:
            await callback.message.edit_text("Ок, запускаю. Дальше буду писать только полезные статусы.")

        fake_message = callback.message
        if fake_message is None:
            return
        await run_user_request(bot, bridge, fake_message, user_id, run.text)

    @dp.message(F.document)
    async def handle_document(message: Message) -> None:
        user_id = message.from_user.id
        if not config.is_allowed(user_id):
            await message.answer(_access_denied_text(user_id))
            return

        try:
            file_name, file_text = await read_document_text(bot, message, config)
        except Exception as exc:
            await message.answer(f"Не смог прочитать файл: {html.escape(str(exc))}")
            return

        caption = (message.caption or "").strip()
        task = caption or "Изучи этот файл и учти его содержимое в работе."
        prompt = (
            f"{task}\n\n"
            f"Файл из Telegram: {file_name}\n"
            "Содержимое файла:\n"
            "```text\n"
            f"{file_text}\n"
            "```"
        )

        if config.confirm_before_run:
            run_id = confirmations.create(user_id, message.chat.id, prompt)
            await message.answer("Файл прочитал. Запустить работу?", reply_markup=_confirm_markup(run_id))
            return

        await run_user_request(bot, bridge, message, user_id, prompt)

    @dp.message(F.text)
    async def handle_text(message: Message) -> None:
        user_id = message.from_user.id
        if not config.is_allowed(user_id):
            await message.answer(_access_denied_text(user_id))
            return

        text = message.text.strip()
        if not text:
            return

        if config.confirm_before_run:
            run_id = confirmations.create(user_id, message.chat.id, text)
            await message.answer("Задачу понял. Запустить?", reply_markup=_confirm_markup(run_id))
            return

        await run_user_request(bot, bridge, message, user_id, text)

    LOG.info("Starting Codex Telegram bridge for %s", config.project_dir)
    await dp.start_polling(bot)


if __name__ == "__main__":
    asyncio.run(main())
